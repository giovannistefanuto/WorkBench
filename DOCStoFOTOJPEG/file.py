import os
from docx import Document

# Parametri principali
k = 3  # Numero di pagine finali da mantenere (puoi cambiarlo qui)
print(f"Numero di pagine da mantenere: {k}")

def trim_docx(input_folder, output_folder, k):
    """Taglia un documento Word (.docx) mantenendo solo le ultime k pagine."""
    # Cerca il file .docx nella cartella di input
    docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]
    if not docx_files:
        raise FileNotFoundError(f"Nessun file DOCX trovato nella cartella '{input_folder}'.")

    # Usa il primo file trovato
    docx_path = os.path.join(input_folder, docx_files[0])
    print(f"Documento trovato: {docx_path}")

    # Carica il documento Word
    doc = Document(docx_path)
    paragraphs = doc.paragraphs

    # Raggruppa i paragrafi in "pagine" virtuali
    lines_per_page = 30
    pages = []
    current_page = []
    for paragraph in paragraphs:
        if len(current_page) >= lines_per_page:
            pages.append(current_page)
            current_page = []
        current_page.append(paragraph.text)
    if current_page:
        pages.append(current_page)

    total_pages = len(pages)

    if k > total_pages:
        raise ValueError(f"Il documento contiene solo {total_pages} pagine, ma k = {k}.")

    # Seleziona le ultime k pagine
    pages_to_keep = pages[-k:]

    # Crea un nuovo documento
    new_doc = Document()

    for page in pages_to_keep:
        for line in page:
            new_doc.add_paragraph(line)
        new_doc.add_paragraph("")  # Aggiungi uno spazio tra le pagine

    # Salva il nuovo documento
    output_path = os.path.join(output_folder, "documento_ridotto.docx")
    new_doc.save(output_path)
    print(f"Documento salvato con successo in '{output_path}'.")

# Esegui la funzione principale
input_folder = r"C:\Users\Utente\Documents\GIOVANNI\PYTHON\WorkBench\DOCStoFOTOJPEG\cartella file"  # Cartella di input
output_folder = r"C:\Users\Utente\Documents\GIOVANNI\PYTHON\WorkBench\DOCStoFOTOJPEG\cartella file"  # Cartella di output

try:
    trim_docx(input_folder, output_folder, k)
    print("Documento ridotto con successo.")
except Exception as e:
    print(f"Errore: {e}")
